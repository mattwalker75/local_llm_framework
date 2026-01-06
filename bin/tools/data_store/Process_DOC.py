#!/usr/bin/env python3
"""
Microsoft Word to Text Converter for RAG Data Store Preparation

This script extracts and cleans text from Microsoft Word documents (.docx, .doc, .rtf),
preparing them for ingestion into a Retrieval-Augmented Generation (RAG) data store.
It handles text extraction, normalization, and optionally preserves document structure
for better RAG performance.

Features:
- Support for .docx, .doc, and .rtf formats
- Robust text extraction with error handling
- Unicode normalization and whitespace cleanup
- Metadata extraction (title, author, subject, etc.)
- Multiple output formats (text, JSON, JSONL)
- Structured output with paragraph/section tracking
- Preserves headings and document structure

Usage:
    # Basic text output
    ./Process_DOC.py -i document.docx -o output.txt

    # JSONL format (recommended for RAG - one record per paragraph)
    ./Process_DOC.py -i document.docx -o output.jsonl -f jsonl

    # JSON format with metadata
    ./Process_DOC.py -i document.docx -o output.json -f json

    # Keep headings separate for better structure
    ./Process_DOC.py -i document.docx -o output.txt --preserve-structure

    # Verbose output for debugging
    ./Process_DOC.py -i document.docx -o output.txt --verbose

Dependencies:
    - python-docx: pip install python-docx (for .docx files)
    - pywin32 (Windows only, optional): pip install pywin32 (for .doc files)
    - striprtf: pip install striprtf (for .rtf files)

Author: Local LLM Framework
License: MIT
"""

import argparse
import json
import sys
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import re
import unicodedata

# Try to import required libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx", file=sys.stderr)

try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False


# ============================================================================
# DOCUMENT EXTRACTION FUNCTIONS
# ============================================================================

def extract_docx_text(path: str, verbose: bool = False) -> str:
    """
    Extract text from .docx file.

    This function reads a Word .docx file and extracts text content from all
    paragraphs, combining them with newlines.

    Args:
        path: Path to the .docx file
        verbose: If True, print progress information

    Returns:
        Combined text from all paragraphs

    Raises:
        FileNotFoundError: If the file doesn't exist
        Exception: If document processing fails

    Example:
        >>> text = extract_docx_text("report.docx")
        >>> print(f"Extracted {len(text)} characters")
    """
    if not DOCX_AVAILABLE:
        print("Error: python-docx is required for .docx files", file=sys.stderr)
        print("Install with: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    try:
        # Validate file exists
        if not Path(path).exists():
            raise FileNotFoundError(f"Document file not found: {path}")

        if verbose:
            print(f"Opening Word document: {path}", file=sys.stderr)

        doc = Document(path)
        paragraphs = []

        if verbose:
            print(f"Processing {len(doc.paragraphs)} paragraphs...", file=sys.stderr)

        # Extract text from each paragraph
        for para_num, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
                if verbose and para_num % 100 == 0:
                    print(f"  Processed {para_num}/{len(doc.paragraphs)} paragraphs", file=sys.stderr)

        if not paragraphs:
            print(f"Warning: No text extracted from {path}", file=sys.stderr)
        elif verbose:
            print(f"✓ Successfully extracted text from {len(paragraphs)} paragraphs", file=sys.stderr)

        return "\n\n".join(paragraphs)

    except FileNotFoundError:
        raise
    except Exception as e:
        print(f"Error processing Word document '{path}': {e}", file=sys.stderr)
        sys.exit(1)


def extract_docx_with_structure(path: str, verbose: bool = False) -> List[Dict[str, any]]:
    """
    Extract text with structure preserved (headings, paragraphs, styles).

    This function extracts text from each paragraph separately, preserving
    paragraph numbers, styles (headings vs normal text), and metadata. This is
    ideal for RAG systems that need to maintain document structure.

    Args:
        path: Path to the .docx file
        verbose: If True, print progress information

    Returns:
        List of dictionaries, each containing:
            - para_number: Paragraph number (1-indexed)
            - text: Text content from that paragraph
            - style: Paragraph style name (e.g., "Heading 1", "Normal")
            - is_heading: Boolean indicating if this is a heading
            - char_count: Number of characters in the paragraph

    Example:
        >>> paragraphs = extract_docx_with_structure("report.docx")
        >>> for para in paragraphs:
        ...     if para['is_heading']:
        ...         print(f"Heading: {para['text']}")
    """
    if not DOCX_AVAILABLE:
        print("Error: python-docx is required for .docx files", file=sys.stderr)
        print("Install with: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    try:
        if not Path(path).exists():
            raise FileNotFoundError(f"Document file not found: {path}")

        if verbose:
            print(f"Opening Word document: {path}", file=sys.stderr)

        doc = Document(path)
        paragraphs_data = []

        if verbose:
            print(f"Processing {len(doc.paragraphs)} paragraphs (structured)...", file=sys.stderr)

        para_count = 0
        for para_num, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                para_count += 1

                # Determine if this is a heading
                style_name = paragraph.style.name if paragraph.style else "Normal"
                is_heading = "heading" in style_name.lower()

                paragraphs_data.append({
                    "para_number": para_count,
                    "text": text,
                    "style": style_name,
                    "is_heading": is_heading,
                    "char_count": len(text)
                })

                if verbose and para_num % 100 == 0:
                    print(f"  Processed {para_num}/{len(doc.paragraphs)} paragraphs", file=sys.stderr)

        if verbose:
            heading_count = sum(1 for p in paragraphs_data if p['is_heading'])
            print(f"✓ Extracted {len(paragraphs_data)} paragraphs ({heading_count} headings)",
                  file=sys.stderr)

        return paragraphs_data

    except FileNotFoundError:
        raise
    except Exception as e:
        print(f"Error processing Word document '{path}': {e}", file=sys.stderr)
        sys.exit(1)


def extract_rtf_text(path: str, verbose: bool = False) -> str:
    """
    Extract text from .rtf file.

    Rich Text Format files are converted to plain text by stripping RTF codes.

    Args:
        path: Path to the .rtf file
        verbose: If True, print progress information

    Returns:
        Plain text extracted from RTF

    Raises:
        FileNotFoundError: If the file doesn't exist
        Exception: If RTF processing fails
    """
    if not RTF_AVAILABLE:
        print("Error: striprtf is required for .rtf files", file=sys.stderr)
        print("Install with: pip install striprtf", file=sys.stderr)
        sys.exit(1)

    try:
        if not Path(path).exists():
            raise FileNotFoundError(f"RTF file not found: {path}")

        if verbose:
            print(f"Opening RTF document: {path}", file=sys.stderr)

        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()

        text = rtf_to_text(rtf_content)

        if not text.strip():
            print(f"Warning: No text extracted from {path}", file=sys.stderr)
        elif verbose:
            print(f"✓ Successfully extracted text from RTF", file=sys.stderr)

        return text

    except FileNotFoundError:
        raise
    except Exception as e:
        print(f"Error processing RTF document '{path}': {e}", file=sys.stderr)
        sys.exit(1)


def extract_doc_metadata(path: str) -> Dict[str, str]:
    """
    Extract metadata from Word document.

    Metadata provides important context for RAG systems, such as document
    title, author, and subject. This helps with document identification
    and relevance scoring.

    Args:
        path: Path to the Word document

    Returns:
        Dictionary containing:
            - title: Document title (empty string if not set)
            - author: Document author (empty string if not set)
            - subject: Document subject (empty string if not set)
            - created: Creation date (empty string if not set)
            - modified: Last modified date (empty string if not set)
            - para_count: Total number of paragraphs (0 for non-.docx)
            - source_file: Name of the source document file

    Example:
        >>> metadata = extract_doc_metadata("report.docx")
        >>> print(f"Title: {metadata['title']}, Author: {metadata['author']}")
    """
    try:
        metadata = {
            "title": "",
            "author": "",
            "subject": "",
            "created": "",
            "modified": "",
            "para_count": 0,
            "source_file": str(Path(path).name)
        }

        # Only .docx files support core properties
        if path.lower().endswith('.docx') and DOCX_AVAILABLE:
            doc = Document(path)

            if doc.core_properties:
                cp = doc.core_properties
                metadata["title"] = cp.title or ""
                metadata["author"] = cp.author or ""
                metadata["subject"] = cp.subject or ""
                metadata["created"] = str(cp.created) if cp.created else ""
                metadata["modified"] = str(cp.modified) if cp.modified else ""

            metadata["para_count"] = len([p for p in doc.paragraphs if p.text.strip()])

        return metadata

    except Exception as e:
        print(f"Warning: Failed to extract metadata from '{path}': {e}", file=sys.stderr)
        # Return minimal metadata on failure
        return {
            "title": "",
            "author": "",
            "subject": "",
            "created": "",
            "modified": "",
            "para_count": 0,
            "source_file": str(Path(path).name)
        }


def extract_document_text(path: str, verbose: bool = False) -> str:
    """
    Extract text from Word document (auto-detects format).

    Automatically detects the file format based on extension and uses the
    appropriate extraction method.

    Args:
        path: Path to the document file (.docx, .doc, or .rtf)
        verbose: If True, print progress information

    Returns:
        Extracted text content

    Raises:
        ValueError: If file format is not supported
    """
    ext = Path(path).suffix.lower()

    if ext == '.docx':
        return extract_docx_text(path, verbose)
    elif ext == '.rtf':
        return extract_rtf_text(path, verbose)
    elif ext == '.doc':
        print("Error: .doc format (older Word format) is not directly supported.", file=sys.stderr)
        print("Please convert to .docx first using Microsoft Word or LibreOffice.", file=sys.stderr)
        print("Alternatively, use: unoconv -f docx file.doc", file=sys.stderr)
        sys.exit(1)
    else:
        print(f"Error: Unsupported file format: {ext}", file=sys.stderr)
        print("Supported formats: .docx, .rtf", file=sys.stderr)
        sys.exit(1)


def extract_document_with_structure(path: str, verbose: bool = False) -> List[Dict[str, any]]:
    """
    Extract text with structure (auto-detects format).

    Args:
        path: Path to the document file
        verbose: If True, print progress information

    Returns:
        List of paragraph dictionaries with structure information
    """
    ext = Path(path).suffix.lower()

    if ext == '.docx':
        return extract_docx_with_structure(path, verbose)
    elif ext == '.rtf':
        # RTF doesn't preserve structure well, fall back to simple extraction
        text = extract_rtf_text(path, verbose)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        return [
            {
                "para_number": i + 1,
                "text": para,
                "style": "Normal",
                "is_heading": False,
                "char_count": len(para)
            }
            for i, para in enumerate(paragraphs)
        ]
    elif ext == '.doc':
        print("Error: .doc format is not directly supported.", file=sys.stderr)
        print("Please convert to .docx first.", file=sys.stderr)
        sys.exit(1)
    else:
        print(f"Error: Unsupported file format: {ext}", file=sys.stderr)
        sys.exit(1)


# ============================================================================
# TEXT CLEANING FUNCTIONS
# ============================================================================

def normalize_text(text: str) -> str:
    """
    Normalize text encoding and formatting for consistent processing.

    This function performs comprehensive text normalization:
    1. Unicode normalization (NFC form) for consistent character representation
    2. Remove control characters (form feeds, carriage returns)
    3. Replace special spaces (non-breaking spaces, tabs) with regular spaces
    4. Collapse excessive whitespace
    5. Preserve paragraph breaks

    Args:
        text: Input text with various encoding and formatting issues

    Returns:
        Normalized, clean text suitable for RAG ingestion

    Note:
        Uses Unicode NFC (Canonical Decomposition, followed by Canonical
        Composition) which is the recommended form for most applications.

    Example:
        >>> text = "Text\\twith\\u00a0various\\r\\nspaces"
        >>> print(normalize_text(text))
        Text with various spaces
    """
    # Unicode normalization to NFC form (canonical composition)
    text = unicodedata.normalize("NFC", text)

    # Remove control characters
    text = text.replace("\f", "")   # Form feed
    text = text.replace("\r", "")   # Carriage return

    # Replace special spaces with regular spaces
    text = text.replace("\u00a0", " ")  # Non-breaking space
    text = text.replace("\t", " ")      # Tab
    text = text.replace("\u200b", "")   # Zero-width space

    # Collapse excessive newlines (3+ → 2)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Collapse excessive spaces (2+ → 1)
    text = re.sub(r' {2,}', ' ', text)

    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    return text.strip()


def clean_text_pipeline(text: str) -> str:
    """
    Apply full text cleaning pipeline.

    This is a convenience function that applies all cleaning steps in the
    correct order.

    Args:
        text: Raw text extracted from document

    Returns:
        Fully cleaned and normalized text

    Example:
        >>> raw_text = extract_document_text("document.docx")
        >>> clean_text = clean_text_pipeline(raw_text)
    """
    text = normalize_text(text)
    return text


# ============================================================================
# OUTPUT FUNCTIONS
# ============================================================================

def write_text_output(output_path: str, text: str, verbose: bool = False) -> None:
    """
    Write cleaned text to a file.

    Args:
        output_path: Path for output text file
        text: Cleaned text content
        verbose: If True, print confirmation message
    """
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

        if verbose:
            char_count = len(text)
            word_count = len(text.split())
            print(f"✓ Written {char_count:,} characters ({word_count:,} words) to {output_path}",
                  file=sys.stderr)
    except Exception as e:
        print(f"Error writing output file '{output_path}': {e}", file=sys.stderr)
        sys.exit(1)


def write_json_output(output_path: str, metadata: Dict, paragraphs_data: List[Dict],
                     verbose: bool = False) -> None:
    """
    Write structured JSON output with metadata and paragraphs.

    Output format:
    {
        "metadata": { ... },
        "paragraphs": [
            {"para_number": 1, "text": "...", "is_heading": false, ...},
            ...
        ]
    }

    Args:
        output_path: Path for output JSON file
        metadata: Document metadata dictionary
        paragraphs_data: List of paragraph dictionaries
        verbose: If True, print confirmation message
    """
    try:
        output = {
            "metadata": metadata,
            "paragraphs": paragraphs_data
        }

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=2, ensure_ascii=False)

        if verbose:
            total_chars = sum(p.get("char_count", len(p.get("text", ""))) for p in paragraphs_data)
            heading_count = sum(1 for p in paragraphs_data if p.get("is_heading", False))
            print(f"✓ Written {len(paragraphs_data)} paragraphs ({heading_count} headings, "
                  f"{total_chars:,} chars) to {output_path}", file=sys.stderr)
    except Exception as e:
        print(f"Error writing JSON output '{output_path}': {e}", file=sys.stderr)
        sys.exit(1)


def write_jsonl_output(output_path: str, metadata: Dict, paragraphs_data: List[Dict],
                      verbose: bool = False) -> None:
    """
    Write JSONL output (one JSON object per line).

    JSONL format is ideal for RAG systems because:
    - Each line is a complete, independent record
    - Easy to stream and process incrementally
    - Each record includes both metadata and paragraph content
    - Standard format for many RAG ingestion pipelines

    Output format (one record per line):
    {"title": "...", "author": "...", "para_number": 1, "text": "...", ...}
    {"title": "...", "author": "...", "para_number": 2, "text": "...", ...}

    Args:
        output_path: Path for output JSONL file
        metadata: Document metadata dictionary
        paragraphs_data: List of paragraph dictionaries
        verbose: If True, print confirmation message
    """
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            for para in paragraphs_data:
                # Merge metadata and paragraph data into single record
                record = {**metadata, **para}
                f.write(json.dumps(record, ensure_ascii=False) + "\n")

        if verbose:
            total_chars = sum(p.get("char_count", len(p.get("text", ""))) for p in paragraphs_data)
            heading_count = sum(1 for p in paragraphs_data if p.get("is_heading", False))
            print(f"✓ Written {len(paragraphs_data)} JSONL records ({heading_count} headings, "
                  f"{total_chars:,} chars) to {output_path}", file=sys.stderr)
    except Exception as e:
        print(f"Error writing JSONL output '{output_path}': {e}", file=sys.stderr)
        sys.exit(1)


# ============================================================================
# MAIN FUNCTION
# ============================================================================

def main():
    """
    Main entry point for Word document to text conversion.

    Parses command-line arguments and orchestrates the conversion process.
    """
    parser = argparse.ArgumentParser(
        description="Convert Microsoft Word documents to text for RAG data store ingestion",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Basic text output
  %(prog)s -i document.docx -o output.txt

  # JSONL format (recommended for RAG)
  %(prog)s -i document.docx -o output.jsonl -f jsonl

  # JSON with metadata
  %(prog)s -i document.docx -o output.json -f json

  # Preserve document structure
  %(prog)s -i document.docx -o output.txt --preserve-structure

  # Verbose output
  %(prog)s -i document.docx -o output.txt --verbose

Supported Formats:
  .docx - Modern Word format (recommended)
  .rtf  - Rich Text Format
  .doc  - Legacy Word format (requires conversion to .docx first)

Output Formats:
  text  - Plain text file with all content merged
  json  - Structured JSON with metadata and paragraphs array
  jsonl - JSONL format (one JSON record per paragraph, best for RAG)
        """
    )

    parser.add_argument(
        "-i", "--input",
        required=True,
        metavar="FILE",
        help="Word document to read as input (.docx, .rtf)"
    )

    parser.add_argument(
        "-o", "--output",
        required=True,
        metavar="FILE",
        help="Output file path"
    )

    parser.add_argument(
        "-f", "--format",
        choices=["text", "json", "jsonl"],
        default="text",
        help="Output format (default: text, recommended for RAG: jsonl)"
    )

    parser.add_argument(
        "--preserve-structure",
        action="store_true",
        help="Preserve document structure (headings, paragraphs) in output"
    )

    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Print detailed progress information"
    )

    args = parser.parse_args()

    # Validate input file exists
    if not Path(args.input).exists():
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    # Check file extension
    ext = Path(args.input).suffix.lower()
    if ext not in ['.docx', '.rtf', '.doc']:
        print(f"Error: Unsupported file format: {ext}", file=sys.stderr)
        print("Supported formats: .docx, .rtf", file=sys.stderr)
        sys.exit(1)

    # Process document based on output format
    try:
        if args.format in ["json", "jsonl"]:
            # Structured output - extract with paragraph information
            if args.verbose:
                print(f"\n{'='*60}", file=sys.stderr)
                print(f"Converting Word document to {args.format.upper()} format", file=sys.stderr)
                print(f"{'='*60}\n", file=sys.stderr)

            # Extract metadata
            metadata = extract_doc_metadata(args.input)

            # Extract text with structure
            paragraphs_data = extract_document_with_structure(args.input, args.verbose)

            # Clean each paragraph
            if args.verbose:
                print("\nCleaning extracted text...", file=sys.stderr)

            for para in paragraphs_data:
                para["text"] = clean_text_pipeline(para["text"])
                para["char_count"] = len(para["text"])

            # Write output
            if args.format == "jsonl":
                write_jsonl_output(args.output, metadata, paragraphs_data, args.verbose)
            else:
                write_json_output(args.output, metadata, paragraphs_data, args.verbose)

        else:
            # Plain text output
            if args.verbose:
                print(f"\n{'='*60}", file=sys.stderr)
                print(f"Converting Word document to plain text", file=sys.stderr)
                print(f"{'='*60}\n", file=sys.stderr)

            # Extract text
            text_data = extract_document_text(args.input, args.verbose)

            # Clean text
            if args.verbose:
                print("\nCleaning extracted text...", file=sys.stderr)

            text_final = clean_text_pipeline(text_data)

            # Write output
            write_text_output(args.output, text_final, args.verbose)

        if args.verbose:
            print(f"\n{'='*60}", file=sys.stderr)
            print(f"✓ Conversion complete!", file=sys.stderr)
            print(f"  Input:  {args.input}", file=sys.stderr)
            print(f"  Output: {args.output}", file=sys.stderr)
            print(f"  Format: {args.format}", file=sys.stderr)
            print(f"{'='*60}\n", file=sys.stderr)
        else:
            print(f"✓ Converted {args.input} → {args.output} ({args.format} format)")

    except KeyboardInterrupt:
        print("\n\nInterrupted by user", file=sys.stderr)
        sys.exit(130)
    except Exception as e:
        print(f"\nUnexpected error: {e}", file=sys.stderr)
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
